"""Form template inspection.

What this module will and will not do:

* **AcroForm PDF** — enumerate real, named fields from the document's own field
  dictionary. Types, options, and required flags come from the file, not guesses.
* **DOCX** — find explicit ``{{placeholder}}`` tokens. Macros are never executed;
  ``.docm`` is rejected outright.
* **Flat PDF** — report page geometry and text only. Field *positions are never
  inferred*: a flat form yields a structured worksheet, and any coordinate mapping
  must be authored and reviewed by a human first.
* **XLSX** — read cells as data. Formulas are returned as their literal text, never
  evaluated.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from typing import Any

from app.forms.enums import (
    FieldDetectionStatus,
    FormFieldType,
    FormFormat,
)
from app.information_registry.enums import Sensitivity

#: DOCX/worksheet placeholder syntax: {{ field_key }}
PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([a-zA-Z0-9_.\-]{1,80})\s*\}\}")

#: Field-name fragments that imply a human must personally execute the field.
_SIGNATURE_HINTS = ("signature", "sign here", "signed by", "/sig", "esign")
_INITIALS_HINTS = ("initial",)
_ATTESTATION_HINTS = ("attest", "certif", "under penalty", "declare", "affirm")
#: Fragments implying sensitive content, used to default a field's sensitivity up.
_SENSITIVE_HINTS = (
    "ssn",
    "social security",
    "ein",
    "tax id",
    "taxpayer",
    "date of birth",
    "dob",
    "account number",
    "routing",
    "passport",
    "driver",
    "license number of individual",
)


class InspectionError(Exception):
    """The template could not be inspected."""


@dataclass(slots=True)
class DetectedField:
    """A field discovered on a template."""

    field_key: str
    label: str
    field_type: str
    native_field_name: str | None = None
    required: bool = False
    allowed_values: list[str] | None = None
    page_number: int | None = None
    max_length: int | None = None
    sensitivity: str = Sensitivity.INTERNAL.value
    instructions: str | None = None
    sort_order: int = 100

    def to_payload(self) -> dict[str, Any]:
        return {
            "field_key": self.field_key,
            "label": self.label,
            "field_type": self.field_type,
            "native_field_name": self.native_field_name,
            "required": self.required,
            "allowed_values": self.allowed_values,
            "page_number": self.page_number,
            "max_length": self.max_length,
            "sensitivity": self.sensitivity,
            "instructions": self.instructions,
            "sort_order": self.sort_order,
        }


@dataclass(slots=True)
class InspectionResult:
    detection_status: str
    form_format: str
    fields: list[DetectedField] = field(default_factory=list)
    page_count: int | None = None
    notes: list[str] = field(default_factory=list)
    #: True when the format cannot be mechanically filled and needs a worksheet.
    worksheet_required: bool = False

    @property
    def field_count(self) -> int:
        return len(self.fields)


def normalise_field_key(name: str, *, used: set[str] | None = None) -> str:
    """Convert a native field name into a stable snake_case key."""
    cleaned = re.sub(r"[^a-zA-Z0-9]+", "_", name or "").strip("_").lower()
    cleaned = re.sub(r"_+", "_", cleaned) or "field"
    if cleaned[0].isdigit():
        cleaned = f"f_{cleaned}"
    cleaned = cleaned[:80]
    if used is None:
        return cleaned
    candidate, index = cleaned, 2
    while candidate in used:
        candidate = f"{cleaned}_{index}"[:80]
        index += 1
    used.add(candidate)
    return candidate


def humanise_label(name: str) -> str:
    """Best-effort human label from a native field name."""
    spaced = re.sub(r"[_\-.]+", " ", name or "").strip()
    spaced = re.sub(r"(?<=[a-z])(?=[A-Z])", " ", spaced)
    return re.sub(r"\s+", " ", spaced).strip().capitalize() or name


def classify_field_type(name: str, native_type: str | None = None) -> str:
    """Infer a field type from the native widget type, then from naming hints.

    Naming hints only ever *escalate* toward human execution (signature, initials,
    attestation). A hint never downgrades a field into something autofillable.
    """
    lowered = (name or "").lower()
    if any(hint in lowered for hint in _SIGNATURE_HINTS):
        return FormFieldType.SIGNATURE.value
    if any(hint in lowered for hint in _INITIALS_HINTS):
        return FormFieldType.INITIALS.value
    if any(hint in lowered for hint in _ATTESTATION_HINTS):
        return FormFieldType.ATTESTATION.value

    if native_type:
        mapping = {
            "/Btn": FormFieldType.CHECKBOX.value,
            "/Ch": FormFieldType.CHOICE.value,
            "/Sig": FormFieldType.SIGNATURE.value,
            "/Tx": FormFieldType.TEXT.value,
        }
        if native_type in mapping:
            return mapping[native_type]

    if "date" in lowered:
        return FormFieldType.DATE.value
    if any(token in lowered for token in ("amount", "fee", "total", "$")):
        return FormFieldType.CURRENCY.value
    if any(token in lowered for token in ("count", "number of", "quantity")):
        return FormFieldType.NUMBER.value
    return FormFieldType.TEXT.value


def infer_sensitivity(name: str, label: str = "") -> str:
    """Default a field's sensitivity upward when its name suggests personal data."""
    haystack = f"{name} {label}".lower()
    if any(hint in haystack for hint in _SENSITIVE_HINTS):
        return Sensitivity.RESTRICTED.value
    if any(token in haystack for token in ("officer", "owner", "control person", "address")):
        return Sensitivity.CONFIDENTIAL.value
    return Sensitivity.INTERNAL.value


def inspect_pdf(content: bytes) -> InspectionResult:
    """Inspect a PDF. Detects AcroForm fields; falls back to flat-PDF reporting."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise InspectionError("pypdf is required to inspect PDF templates.") from exc

    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:
        raise InspectionError(f"The PDF could not be parsed: {type(exc).__name__}") from exc

    if reader.is_encrypted:
        # Attempt the standard empty-password case; anything else needs a human.
        try:
            reader.decrypt("")
        except Exception as exc:
            raise InspectionError("The PDF is encrypted. Supply an unprotected template.") from exc

    page_count = len(reader.pages)
    raw_fields: dict[str, Any] = {}
    try:
        raw_fields = reader.get_fields() or {}
    except Exception:
        # A malformed field tree is not fatal; treat the file as flat.
        raw_fields = {}

    if not raw_fields:
        return InspectionResult(
            detection_status=FieldDetectionStatus.NO_FIELDS_FOUND.value,
            form_format=FormFormat.FLAT_PDF.value,
            page_count=page_count,
            worksheet_required=True,
            notes=[
                "No AcroForm fields were found, so this template is treated as a flat "
                "PDF. Field positions are not inferred; a structured worksheet is "
                "produced instead and any coordinate mapping requires template review.",
            ],
        )

    # Map field name -> page for accurate page numbers.
    page_by_field: dict[str, int] = {}
    for index, page in enumerate(reader.pages, start=1):
        try:
            for annot in page.get("/Annots") or []:
                obj = annot.get_object()
                name = obj.get("/T")
                if name:
                    page_by_field.setdefault(str(name), index)
                parent = obj.get("/Parent")
                if parent is not None:
                    parent_name = parent.get_object().get("/T")
                    if parent_name:
                        page_by_field.setdefault(str(parent_name), index)
        except Exception:
            continue

    used_keys: set[str] = set()
    fields: list[DetectedField] = []
    for order, (native_name, spec) in enumerate(raw_fields.items(), start=1):
        native_type = None
        options: list[str] | None = None
        max_length = None
        required = False
        try:
            native_type = spec.get("/FT")
            raw_options = spec.get("/_States_") or spec.get("/Opt")
            if raw_options:
                options = [str(option) for option in raw_options]
            max_length = spec.get("/MaxLen")
            flags = int(spec.get("/Ff") or 0)
            # Bit 2 (value 2) of the field flags is the "required" flag.
            required = bool(flags & 2)
        except Exception:
            pass

        label = humanise_label(str(native_name))
        field_type = classify_field_type(
            str(native_name), str(native_type) if native_type else None
        )
        fields.append(
            DetectedField(
                field_key=normalise_field_key(str(native_name), used=used_keys),
                label=label,
                field_type=field_type,
                native_field_name=str(native_name),
                required=required,
                allowed_values=options,
                page_number=page_by_field.get(str(native_name)),
                max_length=int(max_length) if max_length else None,
                sensitivity=infer_sensitivity(str(native_name), label),
                sort_order=order * 10,
            )
        )

    signature_count = sum(
        1
        for f in fields
        if f.field_type
        in (
            FormFieldType.SIGNATURE.value,
            FormFieldType.INITIALS.value,
            FormFieldType.ATTESTATION.value,
        )
    )
    notes = [f"Detected {len(fields)} AcroForm field(s) across {page_count} page(s)."]
    if signature_count:
        notes.append(
            f"{signature_count} field(s) require personal execution by an authorised "
            "signatory and will never be auto-filled."
        )
    return InspectionResult(
        detection_status=FieldDetectionStatus.INSPECTED.value,
        form_format=FormFormat.PDF_ACROFORM.value,
        fields=fields,
        page_count=page_count,
        notes=notes,
    )


def inspect_docx(content: bytes) -> InspectionResult:
    """Find ``{{placeholder}}`` tokens in a DOCX template.

    Only explicit placeholders are treated as fields. Guessing at prose would
    produce fields nobody authored, which is precisely the behaviour to avoid.
    """
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise InspectionError("python-docx is required to inspect DOCX templates.") from exc

    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:
        raise InspectionError(f"The DOCX could not be parsed: {type(exc).__name__}") from exc

    texts: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            texts.extend(cell.text for cell in row.cells)
    for section in document.sections:
        for container in (section.header, section.footer):
            if container is not None:
                texts.extend(paragraph.text for paragraph in container.paragraphs)

    used_keys: set[str] = set()
    fields: list[DetectedField] = []
    seen: set[str] = set()
    order = 0
    for text in texts:
        for match in PLACEHOLDER_PATTERN.finditer(text or ""):
            token = match.group(1)
            if token in seen:
                continue
            seen.add(token)
            order += 1
            label = humanise_label(token)
            fields.append(
                DetectedField(
                    field_key=normalise_field_key(token, used=used_keys),
                    label=label,
                    field_type=classify_field_type(token),
                    native_field_name=token,
                    sensitivity=infer_sensitivity(token, label),
                    sort_order=order * 10,
                )
            )

    if not fields:
        return InspectionResult(
            detection_status=FieldDetectionStatus.MANUAL_MAPPING_REQUIRED.value,
            form_format=FormFormat.DOCX.value,
            worksheet_required=True,
            notes=[
                "No {{placeholder}} tokens were found. Add explicit placeholders to the "
                "template, or prepare this form through a worksheet.",
            ],
        )
    return InspectionResult(
        detection_status=FieldDetectionStatus.INSPECTED.value,
        form_format=FormFormat.DOCX.value,
        fields=fields,
        notes=[f"Detected {len(fields)} placeholder token(s)."],
    )


def inspect_xlsx(content: bytes) -> InspectionResult:
    """Report XLSX placeholder cells without evaluating any formula."""
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise InspectionError("openpyxl is required to inspect XLSX templates.") from exc

    try:
        # data_only=False returns the formula *string*; we never compute it.
        workbook = load_workbook(io.BytesIO(content), data_only=False, read_only=True)
    except Exception as exc:
        raise InspectionError(f"The workbook could not be parsed: {type(exc).__name__}") from exc

    used_keys: set[str] = set()
    fields: list[DetectedField] = []
    order = 0
    try:
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if not isinstance(cell.value, str):
                        continue
                    for match in PLACEHOLDER_PATTERN.finditer(cell.value):
                        token = match.group(1)
                        order += 1
                        label = humanise_label(token)
                        fields.append(
                            DetectedField(
                                field_key=normalise_field_key(token, used=used_keys),
                                label=label,
                                field_type=classify_field_type(token),
                                native_field_name=f"{sheet.title}!{cell.coordinate}",
                                sensitivity=infer_sensitivity(token, label),
                                sort_order=order * 10,
                            )
                        )
    finally:
        workbook.close()

    return InspectionResult(
        detection_status=(
            FieldDetectionStatus.INSPECTED.value
            if fields
            else FieldDetectionStatus.MANUAL_MAPPING_REQUIRED.value
        ),
        form_format=FormFormat.XLSX.value,
        fields=fields,
        worksheet_required=not fields,
        notes=[
            "Workbook cells were read as data; no formula was evaluated.",
            f"Detected {len(fields)} placeholder cell(s).",
        ],
    )


def inspect_template(
    content: bytes, *, filename: str, declared_format: str | None = None
) -> InspectionResult:
    """Dispatch inspection by declared format, falling back to the extension."""
    lowered = (filename or "").lower()
    if lowered.endswith((".docm", ".dotm", ".xlsm", ".xltm")):
        raise InspectionError(
            "Macro-enabled Office templates are not accepted. Save the template in a "
            "macro-free format."
        )

    fmt = declared_format
    if fmt in (None, FormFormat.UNKNOWN.value):
        if lowered.endswith(".pdf"):
            fmt = FormFormat.PDF_ACROFORM.value
        elif lowered.endswith((".docx", ".dotx")):
            fmt = FormFormat.DOCX.value
        elif lowered.endswith((".xlsx", ".xltx")):
            fmt = FormFormat.XLSX.value

    if fmt in (FormFormat.PDF_ACROFORM.value, FormFormat.FLAT_PDF.value):
        return inspect_pdf(content)
    if fmt == FormFormat.DOCX.value:
        return inspect_docx(content)
    if fmt == FormFormat.XLSX.value:
        return inspect_xlsx(content)
    if fmt == FormFormat.WEB_WORKSHEET.value:
        return InspectionResult(
            detection_status=FieldDetectionStatus.MANUAL_MAPPING_REQUIRED.value,
            form_format=FormFormat.WEB_WORKSHEET.value,
            worksheet_required=True,
            notes=[
                "A portal form has no downloadable field structure. Declare fields "
                "manually; no browser automation or portal login is performed.",
            ],
        )
    raise InspectionError(f"Unsupported template format for {filename!r}.")


__all__ = [
    "PLACEHOLDER_PATTERN",
    "DetectedField",
    "InspectionError",
    "InspectionResult",
    "classify_field_type",
    "humanise_label",
    "infer_sensitivity",
    "inspect_docx",
    "inspect_pdf",
    "inspect_template",
    "inspect_xlsx",
    "normalise_field_key",
]
