"""Draft form generation.

Guarantees this module upholds:

* The **template is never modified**. Filling always writes a new byte stream.
* A **signature is never fabricated**. Signature, initials, and attestation fields
  are left empty in the output regardless of what a caller passes, and the fact
  that they were skipped is reported.
* **Flattening is opt-in and post-approval only**. A flattened PDF is no longer
  correctable, so it is never the default.
* **Unknown fields are not guessed.** Only fields with an approved mapping and a
  supplied value are written.
"""

from __future__ import annotations

import contextlib
import hashlib
import io
import json
from dataclasses import dataclass, field
from typing import Any

from app.forms.enums import HUMAN_EXECUTION_FIELD_TYPES, FormFieldType, FormFormat
from app.forms.inspection import PLACEHOLDER_PATTERN


class FillingError(Exception):
    """The draft could not be generated."""


@dataclass(frozen=True, slots=True)
class FieldValue:
    """One resolved value destined for a template field."""

    field_key: str
    native_field_name: str | None
    field_type: str
    value: str | None
    #: When true the field is left blank in the output on purpose.
    human_execution_required: bool = False


@dataclass(slots=True)
class FillResult:
    content: bytes
    content_sha256: str
    field_snapshot_sha256: str
    filled_fields: list[str] = field(default_factory=list)
    skipped_signature_fields: list[str] = field(default_factory=list)
    unmatched_fields: list[str] = field(default_factory=list)
    notes: list[str] = field(default_factory=list)
    flattened: bool = False


def field_snapshot_hash(values: list[FieldValue]) -> str:
    """Hash the ordered field set so a draft can be tied to its exact inputs.

    Values are hashed, not stored, so the digest never becomes a channel for
    leaking a restricted value.
    """
    payload = [
        {
            "field_key": item.field_key,
            "value_sha256": (
                hashlib.sha256(item.value.encode("utf-8")).hexdigest() if item.value else None
            ),
            "human_execution_required": item.human_execution_required,
        }
        for item in sorted(values, key=lambda v: v.field_key)
    ]
    canonical = json.dumps(payload, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def _partition(values: list[FieldValue]) -> tuple[list[FieldValue], list[FieldValue]]:
    """Split into writable values and fields a human must execute."""
    writable: list[FieldValue] = []
    human: list[FieldValue] = []
    for item in values:
        if item.human_execution_required or item.field_type in HUMAN_EXECUTION_FIELD_TYPES:
            human.append(item)
        elif item.value is not None and item.value != "":
            writable.append(item)
    return writable, human


def fill_pdf_acroform(
    template: bytes,
    values: list[FieldValue],
    *,
    flatten: bool = False,
    allow_flatten: bool = False,
) -> FillResult:
    """Fill an AcroForm PDF, returning a new document."""
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise FillingError("pypdf is required to fill PDF forms.") from exc

    if flatten and not allow_flatten:
        raise FillingError(
            "Flattening is disabled. A flattened form cannot be corrected, so it is "
            "only permitted as an explicit post-approval step."
        )

    try:
        reader = PdfReader(io.BytesIO(template))
    except Exception as exc:
        raise FillingError(f"The template PDF could not be read: {type(exc).__name__}") from exc

    available = set((reader.get_fields() or {}).keys())
    writable, human = _partition(values)

    writer = PdfWriter(clone_from=reader)
    # Keep field appearances generated so viewers show the values.
    with contextlib.suppress(Exception):
        writer.set_need_appearances_writer(True)

    filled: list[str] = []
    unmatched: list[str] = []
    for item in writable:
        native = item.native_field_name
        if not native or native not in available:
            unmatched.append(item.field_key)
            continue
        text = item.value or ""
        if item.field_type == FormFieldType.CHECKBOX.value:
            # Map truthy text onto the standard AcroForm on/off states.
            text = "/Yes" if str(text).strip().lower() in ("true", "yes", "y", "1", "x") else "/Off"
        for page in writer.pages:
            try:
                writer.update_page_form_field_values(page, {native: text})
            except Exception:
                # A field may not exist on this page; other pages still apply.
                continue
        filled.append(item.field_key)

    if flatten:
        # Remove interactivity only after approval, when corrections are done.
        try:
            for page in writer.pages:
                if "/Annots" in page:
                    del page["/Annots"]
            # pypdf exposes no public API for removing the AcroForm dictionary.
            if "/AcroForm" in writer._root_object:
                del writer._root_object["/AcroForm"]
        except Exception as exc:
            raise FillingError(f"Flattening failed: {type(exc).__name__}") from exc

    buffer = io.BytesIO()
    writer.write(buffer)
    content = buffer.getvalue()

    notes = [f"Filled {len(filled)} field(s) into a new document; the template is unchanged."]
    if human:
        notes.append(
            f"{len(human)} field(s) requiring personal execution were left blank: "
            "no signature, initials, or attestation is ever generated."
        )
    if unmatched:
        notes.append(
            f"{len(unmatched)} mapped value(s) had no matching field in this template "
            "version and were not written."
        )
    return FillResult(
        content=content,
        content_sha256=hashlib.sha256(content).hexdigest(),
        field_snapshot_sha256=field_snapshot_hash(values),
        filled_fields=filled,
        skipped_signature_fields=[item.field_key for item in human],
        unmatched_fields=unmatched,
        notes=notes,
        flattened=flatten,
    )


def fill_docx(template: bytes, values: list[FieldValue]) -> FillResult:
    """Substitute ``{{placeholder}}`` tokens in a DOCX, preserving the template."""
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency is pinned
        raise FillingError("python-docx is required to fill DOCX templates.") from exc

    try:
        document = Document(io.BytesIO(template))
    except Exception as exc:
        raise FillingError(f"The template DOCX could not be read: {type(exc).__name__}") from exc

    writable, human = _partition(values)
    # Signature placeholders are blanked rather than removed, so the printed form
    # still shows where a wet signature belongs.
    replacements = {
        item.native_field_name or item.field_key: (item.value or "") for item in writable
    }
    for item in human:
        replacements.setdefault(item.native_field_name or item.field_key, "")

    filled: set[str] = set()
    key_by_token = {(item.native_field_name or item.field_key): item.field_key for item in values}

    def substitute(text: str) -> str:
        def replace(match: Any) -> str:
            token = match.group(1)
            if token in replacements:
                if token in key_by_token:
                    filled.add(key_by_token[token])
                return str(replacements[token])
            return str(match.group(0))

        return PLACEHOLDER_PATTERN.sub(replace, text)

    def process_paragraph(paragraph: Any) -> None:
        # Substitute run-by-run first to retain per-run formatting; fall back to the
        # whole paragraph when a token is split across runs.
        for run in paragraph.runs:
            if run.text and "{{" in run.text:
                run.text = substitute(run.text)
        if "{{" in paragraph.text and paragraph.runs:
            merged = substitute(paragraph.text)
            if merged != paragraph.text:
                paragraph.runs[0].text = merged
                for run in paragraph.runs[1:]:
                    run.text = ""

    for paragraph in document.paragraphs:
        process_paragraph(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)
    for section in document.sections:
        for container in (section.header, section.footer):
            if container is not None:
                for paragraph in container.paragraphs:
                    process_paragraph(paragraph)

    buffer = io.BytesIO()
    document.save(buffer)
    content = buffer.getvalue()

    unmatched = [item.field_key for item in writable if item.field_key not in filled]
    notes = [f"Substituted {len(filled)} placeholder(s) into a new document."]
    if human:
        notes.append(
            f"{len(human)} signature/attestation placeholder(s) were left blank for "
            "personal execution."
        )
    return FillResult(
        content=content,
        content_sha256=hashlib.sha256(content).hexdigest(),
        field_snapshot_sha256=field_snapshot_hash(values),
        filled_fields=sorted(filled),
        skipped_signature_fields=[item.field_key for item in human],
        unmatched_fields=unmatched,
        notes=notes,
    )


def fill_template(
    template: bytes,
    values: list[FieldValue],
    *,
    form_format: str,
    flatten: bool = False,
    allow_flatten: bool = False,
) -> FillResult:
    """Dispatch to the filler for a fillable format."""
    if form_format == FormFormat.PDF_ACROFORM.value:
        return fill_pdf_acroform(template, values, flatten=flatten, allow_flatten=allow_flatten)
    if form_format == FormFormat.DOCX.value:
        return fill_docx(template, values)
    raise FillingError(
        f"{form_format} cannot be mechanically filled. Generate a field worksheet instead."
    )


def verify_signed_document(
    *,
    approved_draft_sha256: str | None,
    signed_content_sha256: str,
    signed_page_count: int | None = None,
) -> tuple[bool, str]:
    """Sanity-check an uploaded signed copy against the approved draft.

    A signed scan legitimately differs byte-for-byte from the draft, so an exact
    hash match is *not* required — and claiming to verify one would be dishonest.
    What is checked is that the two are not the identical file, since an unchanged
    document means nothing was actually signed.
    """
    if not approved_draft_sha256:
        return False, (
            "No approved draft hash is recorded for this instance, so the signed copy "
            "cannot be related to an approved version."
        )
    if signed_content_sha256 == approved_draft_sha256:
        return False, (
            "The uploaded file is byte-identical to the unsigned approved draft, so it "
            "does not evidence a signature."
        )
    if signed_page_count is not None and signed_page_count <= 0:
        return False, "The uploaded signed document has no readable pages."
    return True, (
        "The signed document differs from the approved draft and is recorded as "
        "signature evidence. Visual confirmation remains a human responsibility."
    )


__all__ = [
    "FieldValue",
    "FillResult",
    "FillingError",
    "field_snapshot_hash",
    "fill_docx",
    "fill_pdf_acroform",
    "fill_template",
    "verify_signed_document",
]
